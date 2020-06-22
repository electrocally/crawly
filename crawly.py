import os
import re
import json
import docx
import click
import fileinput
import pandas as pd
from termcolor import colored, cprint
from pygments import highlight, lexers, formatters

@click.command()
@click.argument('string')
@click.option('--regex', '-r', is_flag=True, help="Search with regex")
@click.option('--case', '-c', is_flag=True, help="Search case sensitive")
@click.option('--replace', '-R', help="String to replace the found strings")
@click.option('--colourless', '-C', is_flag=True, help="Disable colour output")
@click.option('--verbose', '-v', is_flag=True, help="Print more output")
@click.option('--ignore', '-i', help="Directories / Files to ignore")
@click.option('--directory', '-d', default='.', help='The root directory to check')

def main(string, ignore, directory, regex, verbose, case, colourless, replace):

        if verbose and not colourless:
            cprint('Scanned - (Blue) ■', 'blue')
            cprint('Ignored - (Magenta) ■', 'magenta')

        blacklist = [str(ignore)]

        def log(message, colour):
            if verbose:
                if not colourless:
                    cprint(message, ''+colour+'') # this is barbaric.
                elif colourless:
                    print(message)

            elif not verbose and level == 'normal':
                if not colourless:
                    cprint(message, ''+colour+'') # this is *still* barbaric.
                elif colourless:
                    print(message)
                
        def list_build(file, content, line):
            files.append({ 'File' : file, 'Content' : content, "Line": line })

        def check_case(string, line):
            if not regex:
                if case:
                    print(string.lower(), line.lower())

                    return(string.lower(), line.lower())
                else:
                    return(string, line)

        # WIP: Doesn't yet work for case sensitive replacements - No idea how it is for word / excel docs
        def deface(line, thing, replace, full_name):
            try:
                if replace:                    
                    for line in fileinput.FileInput(full_name, inplace=True, backup='.bak'):
                        if case:
                            search = re.compile(thing, re.IGNORECASE)
                            replaced  = search.sub(replace, line)
                            print(line.replace(replaced, replaced), end='')
                        else:
                            print(line.replace(thing, replace), end='')
                            
                    print(f'[+] Replacement in file {full_name}:\n\tBefore: {thing}\n\tAfter: {replace}')
                    
            except Exception as e:
                print(e)

        def extract_from_word(full_name):
            doc = docx.Document(full_name)
            text = []
            ### Headers and Footers
            section = doc.sections[0]
            header = section.header
            for head in header.paragraphs:
                text.append(head.text)
            footer = section.footer
            for foot in footer.paragraphs:
                text.append(foot.text)
                
            ### Paragraphs
            for para in doc.paragraphs:
                text.append(para.text)
            return text

        def extract_from_excel(full_name):
            df = pd.read_excel(full_name, index_col=0)
            return(df)

        def check(content):
            for line in content:
                if len(line) > 0:
                    if regex:
                        reg = re.findall(string, line)
                        for search in reg:
                            if search in line:
                                list_build(full_name, search, line)
                                deface(line, search, replace, full_name)
                    else:
                        check_case(string, line)
                        if string in line:
                            list_build(full_name, string, line)
                            deface(line, string, replace, full_name)


        try:
            level = 'info'
            files = []
            root_dir = directory
            log('\n[+] Discovered Directory Structure:\n', 'blue')
            for dir_name, subdir_list, file_list in os.walk(root_dir):
                log(f'└── {dir_name}', "blue") if dir_name not in blacklist else log(f'└── {dir_name}', "red")
                for fname in file_list:
                    full_name = dir_name+'/'+fname
                    blacklist.append(full_name) if not full_name.endswith(('.txt', '.md', '.yaml', '.json', '.docx', '.xlsx')) else False
                    log(f'  ├── {full_name}', "magenta") if (full_name) in blacklist else log(f'  ├── {full_name}', "blue")
                    
                    if full_name.endswith('.docx') and full_name not in blacklist:
                        if fname.startswith('~$'):
                            cprint(f'[!] Possible Open File: {fname}', 'yellow')
                        check(extract_from_word(full_name))
                    
                    elif full_name.endswith('.xlsx') and full_name not in blacklist:
                         check(extract_from_excel(full_name))

                    elif full_name not in blacklist:
                        with open(full_name) as read_obj:
                            check(read_obj)

            level = 'normal'
            if len(files) != 0:
                if colourless:
                    print(f'\n[+] Discovered Files:\n\n {json.dumps(files, indent=4)}')
                else:
                    print(f'\n[+] Discovered Files:\n\n {highlight(json.dumps(files, indent=4), lexers.JsonLexer(), formatters.TerminalFormatter())}')

            else:
                log('\n[-] No Results', 'yellow')
        
        except Exception as e: 
            level = "normal"
            log(f'[!] Search Failed. Please see error below:\n\t{e}' , 'red')

        except IOError and io:
            level = "normal"
            log(f"[!] Search Failed. Could not open file. Please close running process:\n\t{io}", 'red')

if __name__ == "__main__":
    main()
